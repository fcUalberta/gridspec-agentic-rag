from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
IMAGES = ROOT / "docs" / "images"
OUTPUT = ROOT / "outputs" / "GridSpec_Project_Report.docx"

# standard_business_brief with a named GridSpec brand override.
FONT = "Calibri"
BODY = RGBColor(45, 54, 52)
INK = RGBColor(11, 59, 50)
GREEN = RGBColor(21, 93, 78)
GREEN_MID = RGBColor(43, 122, 99)
GREEN_LIGHT = "EAF5F0"
BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(100, 112, 109)
GRAY_FILL = "F2F4F3"
WHITE = RGBColor(255, 255, 255)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run(run, size=10.5, color=BODY, bold=False, italic=False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Title", 29, INK, 0, 8),
        ("Subtitle", 14, MUTED, 0, 12),
        ("Heading 1", 16, GREEN, 16, 8),
        ("Heading 2", 13, GREEN, 12, 6),
        ("Heading 3", 11.5, BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name, indent, first in (("List Bullet", 0.45, -0.2), ("List Number", 0.45, -0.2)):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(10.5)
        style.font.color.rgb = BODY
        style.paragraph_format.left_indent = Inches(indent)
        style.paragraph_format.first_line_indent = Inches(first)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10

    if "Figure Caption" not in doc.styles:
        cap = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = doc.styles["Figure Caption"]
    cap.font.name = FONT
    cap._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    cap.font.size = Pt(9)
    cap.font.color.rgb = MUTED
    cap.font.italic = True
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(9)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("GRIDSPEC  |  PROJECT REPORT")
    set_run(r, size=8.5, color=MUTED, bold=True)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Engineer-controlled Agentic RAG for RFQ compliance")
    set_run(r, size=8, color=MUTED)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run(run, size=8, color=MUTED)


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text.upper())
    set_run(r, size=9, color=GREEN_MID, bold=True)


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.75)
    r = p.add_run(text)
    set_run(r)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run(r)


def add_callout(doc: Document, label: str, text: str, fill=GREEN_LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    set_run(r, size=10.5, color=INK, bold=True)
    r = p.add_run(text)
    set_run(r, size=10.5, color=BODY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN_LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(header)
        set_run(r, size=9.3, color=INK, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            if row_index % 2:
                set_cell_shading(cells[idx], "F8FAF9")
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(value)
            set_run(r, size=9.1)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc: Document, filename: str, caption: str, width=6.35) -> None:
    path = IMAGES / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph(style="Figure Caption")
    cp.add_run(caption)


def page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_prompt(doc: Document, title: str, prompt: str, outcome: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run(r, size=11.5, color=GREEN, bold=True)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F9F8")
    p = cell.paragraphs[0]
    r = p.add_run(f'"{prompt}"')
    set_run(r, size=9.7, color=BODY, italic=True)
    p = cell.add_paragraph()
    r = p.add_run("Result: ")
    set_run(r, size=9.4, color=INK, bold=True)
    r = p.add_run(outcome)
    set_run(r, size=9.4)


def build() -> None:
    doc = Document()
    style_document(doc)
    for section in doc.sections:
        add_header_footer(section)
        add_page_number(section)

    # Cover page: editorial_cover pattern with GridSpec brand override.
    cover = doc.sections[0]
    cover.different_first_page_header_footer = True
    cover.first_page_header.paragraphs[0].text = ""
    cover.first_page_footer.paragraphs[0].text = ""
    for _ in range(4):
        doc.add_paragraph()
    add_kicker(doc, "Project report")
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("GridSpec Agentic RAG for RFQ Compliance")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Engineer-controlled requirements extraction, evidence-backed compliance, cohesive solution design, and BOQ/BOM preparation")
    doc.add_paragraph()
    add_callout(
        doc,
        "Value proposition",
        "GridSpec helps tendering and proposal teams replace several days of manual RFQ and product-catalog review with an initial bid draft targeted for completion in under one hour. Every compliance outcome is designed to be linked to verifiable evidence or explicitly routed for clarification.",
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Prepared from the implemented GridSpec prototype and its completed demonstration workflow")
    set_run(r, size=10, color=MUTED, italic=True)
    p = doc.add_paragraph()
    r = p.add_run("August 2026")
    set_run(r, size=10, color=MUTED, bold=True)
    page_break(doc)

    add_kicker(doc, "Executive overview")
    doc.add_heading("Project overview", level=1)
    add_body(
        doc,
        "GridSpec is a controlled Agentic RAG application for engineering tender and proposal teams. It converts customer RFQs and approved product documentation into a structured requirement register, an evidence-backed compliance matrix, a cohesive technical solution, and a consolidated BOQ/BOM.",
    )
    add_body(
        doc,
        "The system deliberately avoids a free-running multi-agent design. Deterministic code handles predictable processing, retrieval-augmented generation supplies bounded product evidence, reasoning agents are introduced only for ambiguous interpretation and engineering judgment, and humans review the major stage outputs.",
    )
    add_callout(
        doc,
        "Success target",
        "Produce an initial, reviewable bid draft in under one hour, with every compliance decision either linked to validated evidence or explicitly routed to an engineer for clarification.",
    )
    doc.add_heading("What the project delivers", level=2)
    for item in (
        "A page-traceable requirement register extracted from real PDF specifications.",
        "Requirement-to-product matching grounded in indexed product-manual evidence.",
        "Conservative compliance decisions with citations, alternatives, and clear unknowns.",
        "A cross-requirement cohesion review covering interfaces, assumptions, and deviations.",
        "A consolidated BOQ/BOM plus CSV and JSON proposal outputs.",
        "Checkpointed jobs that can resume after interruption without repeating completed work.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Problem statement", level=1)
    add_body(
        doc,
        "Engineering RFQs can span hundreds of pages and mix product capabilities, system architecture, testing, documentation, delivery, bidder qualification, and owner responsibilities. Tendering teams must manually locate obligations, interpret them consistently, find suitable evidence, document gaps, propose deviations, and assemble a coherent offer.",
    )
    add_body(
        doc,
        "The manual workflow typically takes several days and is difficult to audit. A generic LLM-only workflow can be faster, but may omit requirements, confuse products, overlook numeric or standards mismatches, or generate unsupported positive claims. GridSpec addresses both problems with a controlled workflow that preserves source traceability and uses agents only after deterministic checks.",
    )
    doc.add_heading("Project objectives", level=2)
    for item in (
        "Reduce the time required to create an initial RFQ response draft.",
        "Keep requirements and compliance outcomes traceable to exact source pages and quotations.",
        "Reduce unnecessary model calls through deterministic extraction, routing, and filtering.",
        "Prevent positive compliance conclusions when product evidence is missing or unverifiable.",
        "Keep engineers in control without requiring approval of every individual requirement.",
        "Produce a technically cohesive offer rather than a disconnected set of product matches.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("Solution overview", level=1)
    add_figure(doc, "solution-overview.png", "Figure 1. GridSpec Agentic RAG solution overview")
    add_body(
        doc,
        "The flow begins with the customer RFQ and product manuals, converts the RFQ into a reviewable requirement story, retrieves product evidence, evaluates compliance, checks the overall solution, and exports the resulting proposal package.",
    )
    for label, detail in (
        ("Inputs", "Customer RFQ, product manuals, and supporting technical sources."),
        ("Requirements", "Atomic obligations organized into solution packages and engineering subcategories, with an engineer edit checkpoint."),
        ("Agentic RAG compliance", "Evidence retrieval, hard-value filtering, bounded evaluation, and citation validation."),
        ("Cohesive solution", "Cross-requirement compatibility, interfaces, alternatives, deviations, and assumptions."),
        ("Outputs", "Compliance matrix, recommended solution, BOQ/BOM, and machine-readable exports."),
    ):
        add_body(doc, f"{label}: {detail}", bold_lead=f"{label}:")

    page_break(doc)
    doc.add_heading("Design principles", level=2)
    principles = [
        ("Deterministic first", "Rules, filters, validation, persistence, and aggregation are implemented in code."),
        ("Evidence before reasoning", "Product compliance is evaluated only after suitable manual evidence has been retrieved."),
        ("Agents only where necessary", "Agents are limited to ambiguous interpretation, grounded compliance reasoning, and solution cohesion."),
        ("No unsupported positive claims", "A positive product decision requires validated evidence."),
        ("Human in the loop", "Requirements and the complete solution remain reviewable and editable."),
        ("Resumable by design", "Long extraction and compliance jobs persist checkpoints and resume incomplete batches."),
    ]
    add_table(doc, ["Principle", "Application"], [[a, b] for a, b in principles], [2600, 6760])

    page_break(doc)
    doc.add_heading("Datasets and source material", level=1)
    add_body(
        doc,
        "The prototype was exercised with one public utility tender document and two product manuals. The application stores the original PDFs locally, derives page-aware text and citations, and creates a vector evidence index from product-manual chunks. The source documents themselves are not included in this report.",
    )
    dataset_rows = [
        [
            "Customer RFQ",
            "20200820Draft_Tor_Bid_No._TS12-DSS-02.pdf",
            "339 pages; 7.8 MB",
            "EGAT draft invitation to bid for a Digital Substation Protection and Automation System, Transmission System Expansion Project No. 12.",
        ],
        [
            "Product manual",
            "B30-1601-0109-87x-2.pdf",
            "825 pages; 25.7 MB",
            "GE Vernova UR Family B30 Bus Differential System instruction manual, product version 8.71.",
        ],
        [
            "Product manual",
            "URCG-1601-0401-87x.pdf",
            "604 pages; 8.9 MB",
            "GE Vernova UR Family Communications Guide, product version 8.7x.",
        ],
        [
            "Derived data",
            "Requirement and compliance records",
            "1,110 requirements",
            "Normalized requirements, engineering taxonomy, evidence routes, compliance outcomes, citations, job checkpoints, and solution context.",
        ],
    ]
    add_table(doc, ["Type", "Source", "Scale", "Use in workflow"], dataset_rows, [1350, 2500, 1300, 4210])
    add_callout(
        doc,
        "Data-use note",
        "The product manuals are copyrighted vendor documentation and are used locally as evidence sources. Any redistribution or public use of the original documents must follow the document owners' terms. The demonstration metrics are indicative workflow results, not an independently adjudicated accuracy benchmark.",
        fill="FFF8E5",
    )
    doc.add_heading("Derived stores", level=2)
    for item in (
        "Page-aware RFQ blocks with document ID, page number, section context, quotation, and bounding box.",
        "Normalized requirements with solution package, subcategory, lifecycle phase, evidence scope, and expected evidence.",
        "Qdrant product-evidence vectors containing manual text, product document identity, and page metadata.",
        "SQLite operational state for documents, requirements, assessments, solutions, jobs, batches, and audit-oriented metadata.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("Key capabilities", level=1)
    capabilities = (
        "Real PDF upload and page-aware parsing with PyMuPDF.",
        "Selective complex-layout processing with LiteParse.",
        "Exact source quotations, page numbers, and bounding-box metadata.",
        "Deterministic obligation detection and section-context filtering.",
        "Clause reassembly, text repair, and numeric/standards-aware deduplication.",
        "Engineering taxonomy across products, system design, services, testing, and deliverables.",
        "Editable requirement register with an engineer checkpoint.",
        "Product-manual chunking, embeddings, and local Qdrant indexing.",
        "Hybrid semantic and lexical evidence matching.",
        "Hard gates for weak, numerically inconsistent, or wrong-product evidence.",
        "Ollama-first compliance evaluation with selective Fireworks escalation.",
        "Citation and evidence-quote validation.",
        "Checkpointed compliance batches with resume and re-evaluate controls.",
        "Cross-requirement cohesion checks and alternative/deviation handling.",
        "CSV compliance export and JSON solution export.",
    )
    for item in capabilities:
        add_bullet(doc, item)

    doc.add_heading("Requirement organization", level=2)
    add_body(
        doc,
        "Requirements are ordered as a solution narrative instead of a flat list. Top-level solution packages include design basis and standards, protection and control, process bus and instrument interfaces, station communications, HMI and automation, panels and auxiliary systems, metering, engineering services, verification, and generic items. Subcategories then provide the evidence route and reading order within each package.",
    )

    page_break(doc)
    doc.add_heading("Architecture overview", level=1)
    add_figure(doc, "architecture-overview.png", "Figure 2. GridSpec end-to-end Agentic RAG architecture")
    add_body(
        doc,
        "The architecture separates the user experience, orchestration, deterministic document processing, agentic RAG compliance, controlled tools, model routing, persistence, and final proposal assembly.",
    )
    architecture_points = [
        ("User experience", "A React workspace guides the engineer through Sources, Requirements, Compliance, Solution, and Outputs."),
        ("Controlled orchestration", "FastAPI exposes the application API while LangGraph coordinates explicit state transitions, background jobs, and checkpoints."),
        ("Document processing", "Local parsers preserve source traceability. Deterministic candidate and taxonomy logic handle predictable transformations."),
        ("Agentic RAG", "The workflow scopes the query, retrieves product evidence through MCP, applies hard gates, evaluates bounded evidence, and validates citations."),
        ("Persistence", "SQLite stores operational state and Qdrant stores source-aware product evidence vectors."),
    ]
    for label, detail in architecture_points:
        add_body(doc, f"{label}: {detail}", bold_lead=f"{label}:")

    doc.add_heading("RAG is a technique, not an agent", level=2)
    add_body(
        doc,
        "The RAG path performs query construction, evidence-scope routing, vector and lexical retrieval, hard-value filtering, context bounding, grounded evaluation, and citation validation. Qdrant finds candidate evidence; it does not determine compliance.",
    )
    doc.add_heading("MCP is the controlled tool boundary", level=2)
    add_body(
        doc,
        "The Product Catalog MCP server exposes a limited set of indexing, search, batch retrieval, and status operations. Reasoning nodes do not receive unrestricted database or file-system access.",
    )
    doc.add_heading("Controlled model routing", level=2)
    for item in (
        "Ollama is the first-choice local model for ambiguous requirement candidates and evidence-qualified compliance evaluation.",
        "Fireworks AI provides embeddings and is reserved for fallback or escalation when cases are low-confidence, conflicting, or consequential.",
        "Structured Pydantic schemas constrain model output and timeouts prevent indefinite requests.",
        "Provider failures and low-confidence results are persisted conservatively rather than converted into unsupported positive claims.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("End-to-end decision path", level=1)
    add_figure(doc, "decision-path.png", "Figure 3. GridSpec end-to-end decision path")
    add_body(
        doc,
        "For every requirement, the system determines the evidence scope, calls approved MCP functions, retrieves and filters candidate evidence, evaluates only qualified evidence, validates citations, and records the result. Requirements without sufficient evidence are intentionally routed to Unknown, clarification, or manual review.",
    )
    decision_steps = (
        "Parse documents and preserve page-level citations.",
        "Extract, normalize, classify, and review requirements.",
        "Build a requirement query and select the appropriate evidence scope.",
        "Retrieve and rank evidence through MCP and RAG.",
        "Apply vector, lexical, numeric, standards, and model-family gates.",
        "Invoke the compliance agent with bounded evidence only.",
        "Validate citations before persisting a positive decision.",
        "Repeat through checkpointed batches until all requirements are processed.",
        "Use the cohesion agent to identify interfaces, conflicts, and alternatives.",
        "Build the consolidated BOQ/BOM and send the complete solution to engineer review.",
    )
    for step in decision_steps:
        add_number(doc, step)

    page_break(doc)
    doc.add_heading("Application workflow", level=1)
    doc.add_heading("1. Sources", level=2)
    add_body(
        doc,
        "The Sources workspace accepts RFQ and product-manual PDFs, shows parsing and indexing state, and reports resumable extraction progress. Product manuals are chunked, embedded, and indexed through the MCP service; RFQs remain the source of customer obligations.",
    )
    add_figure(doc, "app-sources.png", "Figure 4. Sources workspace with one RFQ, two product manuals, and completed extraction")

    page_break(doc)
    doc.add_heading("2. Requirements", level=2)
    add_body(
        doc,
        "The Requirements workspace presents an ordered engineering story by solution package and subcategory. Every item retains its RFQ page reference and expected evidence route. Editing is optional, and a single stage-level action continues the workflow without requiring individual approvals.",
    )
    add_figure(doc, "app-requirements.png", "Figure 5. Engineer requirement-review checkpoint")

    page_break(doc)
    doc.add_heading("3. Compliance", level=2)
    add_body(
        doc,
        "The Compliance workspace separates product claims from system-design, deliverable, testing, and engineer-review obligations. It reports checkpoint progress, model routing, result distribution, batch failures, and expandable evidence. Failed or interrupted batches can resume without repeating completed work.",
    )
    add_figure(doc, "app-compliance.png", "Figure 6. Completed controlled compliance run and evidence lanes")

    page_break(doc)
    doc.add_heading("4. Cohesive solution", level=2)
    add_body(
        doc,
        "The Solution workspace assembles the recommended protection-panel configuration and places the BOM beside interface and cohesion findings. The model receives a compact, grouped context so it can reason across the full matrix without hiding unresolved work.",
    )
    add_figure(doc, "app-solution-bom-cohesion.png", "Figure 7. Generated solution, BOM, and interface checks")

    page_break(doc)
    doc.add_heading("5. Assumptions and deviations", level=2)
    add_body(
        doc,
        "Assumptions and deviations are preserved as first-class proposal artifacts. This makes unsupported dependencies, required third-party evidence, exceptions, and engineering decisions visible before bid submission.",
    )
    add_figure(doc, "app-solution-assumptions-deviations.png", "Figure 8. Solution assumptions and deviation register")

    page_break(doc)
    doc.add_heading("6. Outputs", level=2)
    add_body(
        doc,
        "The Outputs workspace produces a CSV compliance matrix and a JSON solution package from persisted live results. The CSV includes requirement context, decision, product, rationale, evidence, and confidence. The JSON contains the BOM, cohesion checks, assumptions, and deviations.",
    )
    add_figure(doc, "app-outputs.png", "Figure 9. Exportable compliance matrix and solution package")

    page_break(doc)
    doc.add_heading("Technology stack", level=1)
    stack_rows = [
        ["Web UI", "React 19, TypeScript, Vinext, Vite", "Five-stage engineer workspace, editing, progress, review, and exports"],
        ["API", "Python, FastAPI, Uvicorn", "Uploads, background jobs, status, persistence access, and workflow endpoints"],
        ["Workflow", "LangGraph, LangChain", "Explicit extraction, compliance, and solution graphs"],
        ["PDF", "PyMuPDF, LiteParse", "Canonical extraction and selective complex-layout processing"],
        ["Local model", "Ollama", "Ambiguous candidate interpretation and primary compliance evaluation"],
        ["Hosted AI", "Fireworks AI", "Embeddings, fallback reasoning, and selective escalation"],
        ["Tools", "Model Context Protocol", "Controlled catalog indexing and evidence retrieval"],
        ["Vector store", "Qdrant local mode", "Product-manual vectors, evidence text, and source metadata"],
        ["State", "SQLite", "Documents, requirements, jobs, batches, decisions, and solutions"],
        ["Schemas", "Pydantic", "Structured API, agent, and workflow outputs"],
        ["Quality", "Pytest, Ruff, ESLint", "Behavioral tests and static quality checks"],
    ]
    add_table(doc, ["Layer", "Technology", "Responsibility"], stack_rows, [1500, 2750, 5110])

    page_break(doc)
    doc.add_heading("API and MCP interfaces", level=1)
    add_body(
        doc,
        "The API exposes health and status, document upload, extraction jobs, requirement review, compliance jobs, and solution generation. The MCP server keeps product evidence access bounded to four tools.",
    )
    mcp_rows = [
        ["index_product_manual", "Embed and index page-aware product-manual chunks."],
        ["search_manual_evidence", "Retrieve source-aware evidence for one query."],
        ["search_manual_evidence_batch", "Retrieve evidence for a checkpoint batch."],
        ["catalog_status", "Report catalog readiness and indexed chunk count."],
    ]
    add_table(doc, ["MCP tool", "Purpose"], mcp_rows, [3100, 6260])
    api_rows = [
        ["Documents", "List and upload RFQ or product PDFs"],
        ["Extraction", "Start, force, resume, and inspect requirement extraction"],
        ["Requirements", "List and edit normalized requirements"],
        ["Compliance", "Start, resume, re-evaluate, and inspect checkpointed decisions"],
        ["Solution", "Generate and retrieve the cohesive solution"],
    ]
    add_table(doc, ["API area", "Responsibility"], api_rows, [2300, 7060])

    page_break(doc)
    doc.add_heading("Prompts used during vibe coding", level=1)
    add_body(
        doc,
        "The following are representative user-facing development prompts that shaped the implementation. They are lightly edited for clarity and do not include hidden system instructions or private reasoning.",
    )
    prompts = [
        (
            "1. Real pipeline, not a mock",
            "I need real processing and a real model connection. I do not want dummy data or dummy pipelines. Even if it is one sample, new data should work when I add it.",
            "Replaced prototype-only behavior with FastAPI, persisted records, real PDF parsing, model calls, and live UI state.",
        ),
        (
            "2. Python-first controlled architecture",
            "I want it Python based, LangGraph and LangChain based, and as open as possible. Introduce an agent only when it is necessary.",
            "Moved the core workflow into Python and represented extraction, compliance, and solution assembly as explicit controlled graphs.",
        ),
        (
            "3. Hosted and local model routing",
            "I have Fireworks AI access. Can I go without local models, and can we reduce Fireworks API calls?",
            "Added Fireworks configuration and later introduced Ollama-first evaluation with selective Fireworks escalation.",
        ),
        (
            "4. Engineer-centered UI",
            "The UI should accept inputs, show results, and get engineer feedback at requirement extraction, compliance analysis, and major output stages.",
            "Created five workflow pages with stage checkpoints, editable requirements, progress, result filters, solution review, and exports.",
        ),
        (
            "5. Faster extraction",
            "Extraction is taking too long. Instead of asking an LLM to parse the whole document, what other methods can provide similar results?",
            "Shifted to PyMuPDF, selective LiteParse, deterministic candidate rules, page triage, source validation, and model calls only for ambiguous candidates.",
        ),
        (
            "6. Controlled compliance",
            "Use deterministic pre-matching, Ollama-first evaluation, Fireworks escalation only when necessary, resumable batches, and stronger qualification and filtering.",
            "Produced the controlled compliance pipeline with hard gates, bounded evidence, structured decisions, checkpoints, and conservative fallbacks.",
        ),
        (
            "7. Cohesive requirement story",
            "Order requirements as Products, System Design, and Services and Deliverables, with subcategories, so the result reads cohesively.",
            "Introduced solution packages, engineering chapters, evidence scopes, and deterministic ordering across extraction, compliance, and solution assembly.",
        ),
        (
            "8. Compliance usability and recovery",
            "Compress the compliance view, show the distribution, and add resume and re-evaluate controls for failed batches.",
            "Added summary metrics, evidence lanes, pagination, compact rows, progress counters, resume, and force re-evaluation.",
        ),
    ]
    for title, prompt, outcome in prompts:
        add_prompt(doc, title, prompt, outcome)

    doc.add_heading("Runtime prompt patterns", level=2)
    add_body(
        doc,
        "Production prompts use structured outputs and conservative wording. The full prompt text is implemented in the source code; the patterns below summarize their intended behavior.",
    )
    runtime_prompts = [
        ("Ambiguous requirement interpretation", "Accept only supplier-facing technical obligations. Reject legends, owner actions, bidder administration, prices, and descriptive background. Use source IDs supplied by the deterministic parser and do not invent citations."),
        ("Controlled compliance evaluation", "Decide only from supplied evidence candidates. Require explicit support for every material element. Use only provided evidence IDs, keep the rationale technical, and request escalation only for real conflicts or consequential uncertainty."),
        ("Cohesive solution assembly", "Assemble one solution from the compliance summary without introducing unverified capabilities. Check protection independence, CT/VT interfaces, DC burden, communications, synchronization, test facilities, environmental design, panel construction, and engineering scope."),
    ]
    for label, text in runtime_prompts:
        add_body(doc, f"{label}: {text}", bold_lead=f"{label}:")

    page_break(doc)
    doc.add_heading("Iterations tried", level=1)
    iterations = [
        ("Interactive concept and UI prototype", "Established the five-stage workflow and engineer checkpoints. The early implementation demonstrated the experience but was not sufficient as a data-processing system."),
        ("Real Python pipeline", "Introduced FastAPI, LangGraph, LangChain, SQLite, Fireworks, and a real upload flow. Unused TypeScript server files were removed as the Python backend became authoritative."),
        ("LLM-heavy extraction", "Early extraction sent too much document content through model calls. It was slow, appeared stuck at zero percent, and was vulnerable to provider latency and structured-output failures."),
        ("Controlled hybrid extraction", "PyMuPDF became the canonical parser; LiteParse was restricted to difficult pages; explicit obligations were extracted deterministically; ambiguous candidates alone went to a model."),
        ("Progress and resilience", "Extraction was moved into dynamic background batches with page triage, status reporting, retry, source validation, and restart-safe checkpoints."),
        ("Flat to cohesive taxonomy", "Requirements were reorganized into solution packages and subcategories, improving both the reading order and the evidence route used by compliance."),
        ("Model-per-requirement compliance", "Direct API evaluation proved slow and produced too many unknown or weakly supported outcomes because retrieval quality and requirement type were not sufficiently controlled."),
        ("Agentic RAG compliance", "Added MCP batch retrieval, Qdrant, lexical overlap, hard-value gates, evidence scopes, Ollama-first evaluation, selective Fireworks escalation, and citation validation."),
        ("Resumable compliance experience", "Added eight-requirement checkpoints, resume-failed-only behavior, re-evaluation, visible progress, result distribution, evidence lanes, and a compact paginated table."),
        ("Solution cohesion and outputs", "Added cross-requirement solution checks, alternates, deviations, assumptions, BOM construction, and CSV/JSON exports."),
    ]
    for idx, (name, detail) in enumerate(iterations, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{idx}. {name}")
        set_run(r, size=11.2, color=GREEN, bold=True)
        add_body(doc, detail)

    page_break(doc)
    doc.add_heading("Demonstration snapshot", level=1)
    add_body(
        doc,
        "The screenshots supplied for this report show one completed run. Fireworks escalations are a subset of the locally evaluated cases, not an additional requirement count.",
    )
    metric_rows = [
        ["RFQ pages", "339", "Input document scale"],
        ["Extraction batches", "21 of 21", "Completed"],
        ["Extracted requirements", "1,110", "Candidate requirements after processing"],
        ["Compliance decisions", "1,110", "100 percent processed"],
        ["Deterministic or routed", "574", "No product-model evaluation required"],
        ["Ollama evaluations", "536", "Primary model route"],
        ["Fireworks escalations", "112", "Subset escalated for difficult cases"],
        ["Product compliant", "214", "Evidence-backed positive results"],
        ["Product exceptions", "61", "Non-compliant or conditional product cases"],
        ["Product unresolved", "556", "Evidence insufficient or unknown"],
        ["Engineering evidence due", "279", "System design, deliverables, tests, or engineer confirmation"],
        ["Batch failures", "0", "Completed screenshot state"],
    ]
    add_table(doc, ["Metric", "Observed value", "Interpretation"], metric_rows, [2500, 1700, 5160])
    add_callout(
        doc,
        "Evaluation caveat",
        "These counts demonstrate workflow completion and routing behavior. They do not establish extraction precision, recall, or engineering decision accuracy because an independently labeled gold-standard dataset has not yet been created.",
        fill="FFF8E5",
    )

    page_break(doc)
    doc.add_heading("Learnings and observations", level=1)
    learnings = [
        ("Vector retrieval is necessary but not sufficient", "Semantic similarity is useful for finding evidence, but it cannot prove numeric ratings, standards, variants, or product-family applicability. Hard gates remain necessary."),
        ("The parser must own citations", "Page numbers and quotations should come from the local parser. A model can select an evidence ID, but it should not invent the evidence text or source location."),
        ("Unknown is a valid engineering outcome", "Forcing every requirement into Compliant or Non-compliant creates false certainty. Missing evidence should remain visible as Unknown, clarification required, or engineering evidence due."),
        ("Requirement type determines the compliance route", "System design, testing, deliverables, and bid commitments cannot be proved from a product manual. Evidence-scope routing reduces meaningless searches and misleading unknowns."),
        ("Taxonomy improves more than presentation", "A cohesive package and subcategory structure improves reading order, retrieval queries, compliance grouping, solution context, and engineer review."),
        ("Deterministic processing reduces latency and cost", "Rules can identify explicit obligations, reject procurement noise, normalize text, route evidence, and reject weak matches without a model call."),
        ("Local-first routing needs bounded workloads", "Small local models are useful when prompts and schemas are narrow. They become slow or inconsistent when asked to process large documents or unconstrained evidence sets."),
        ("Checkpointing is a product feature", "For large RFQs, resume and retry are not merely backend concerns. Visible progress and recovery directly affect user trust."),
        ("Stage-level review is more usable than per-item approval", "Engineers need the ability to edit any requirement, but requiring 1,110 individual approvals would make the workflow unusable."),
        ("Solution-level reasoning is distinct from compliance", "Individually compliant products can still form an incoherent system. Interfaces, dependencies, quantities, and assumptions require a second, bounded reasoning stage."),
    ]
    for title, detail in learnings:
        add_body(doc, f"{title}: {detail}", bold_lead=f"{title}:")

    doc.add_heading("Quality validation", level=1)
    add_body(
        doc,
        "The repository was validated after documentation updates. The backend test suite completed with 38 passing tests. Ruff, ESLint, and the frontend production build also completed successfully. Tests cover API behavior, extraction jobs and resilience, taxonomy, compliance checkpoints, and bounded solution context.",
    )

    doc.add_heading("Security and responsible use", level=1)
    for item in (
        "Environment files, API keys, uploaded PDFs, SQLite files, and Qdrant runtime data must remain outside version control.",
        "The MCP server should remain a bounded interface with approved functions rather than unrestricted storage access.",
        "Uploaded source documents should follow organizational retention, confidentiality, and intellectual-property policies.",
        "Authentication, authorization, encrypted storage, managed secrets, and tenant isolation are required before multi-user deployment.",
        "Every bid response must be reviewed by qualified protection, control, panel, application, and commercial engineers.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Current limitations", level=1)
    for item in (
        "Only PDF source documents are supported.",
        "Scanned and highly graphical specifications can still require manual review.",
        "Compliance quality is limited by the completeness and correctness of uploaded product evidence.",
        "The demonstration dataset does not yet include independent ground-truth labels.",
        "SQLite and local Qdrant are suited to a prototype or single-node workflow, not horizontal scale.",
        "The BOQ/BOM remains a proposal aid and requires engineering and commercial validation.",
        "The current local application does not implement authentication or role-based access control.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("Future work", level=1)
    future = [
        ("Gold-standard evaluation set", "Create a manually adjudicated RFQ and product-evidence corpus. Measure requirement extraction precision and recall, evidence retrieval recall at k, citation validity, decision accuracy, and engineer acceptance."),
        ("Advanced document intelligence", "Add robust OCR for image-only pages, table reconstruction, drawing and schedule extraction, multilingual handling, and section-aware document navigation."),
        ("Retrieval and reranking", "Introduce BM25 or another indexed lexical layer, a technical reranker, better product-variant metadata, query expansion, and configurable evidence policies by requirement type."),
        ("Structured product catalog", "Represent models, order codes, options, accessories, ratings, compatibility, and lifecycle status as structured entities rather than relying entirely on manual text."),
        ("Engineer feedback learning", "Capture edits, accepted decisions, rejected evidence, and preferred alternatives to improve rules, thresholds, prompts, and evaluation sets."),
        ("Richer proposal outputs", "Generate customer-ready Word and PDF responses, Excel compliance matrices, deviation schedules, clause-by-clause responses, and traceable BOQ workbooks."),
        ("Enterprise integration", "Connect product lifecycle, document management, CRM, CPQ, ERP, standards libraries, and pricing systems through controlled MCP tools."),
        ("Production architecture", "Move operational state to PostgreSQL, use a managed vector service or Qdrant server, store documents in object storage, and add queues, workers, observability, backups, and disaster recovery."),
        ("Security and governance", "Add single sign-on, role-based access, tenant isolation, encryption, audit exports, model and prompt versioning, data retention, and approval policies."),
        ("Performance and cost controls", "Add per-stage latency and token metrics, embedding caches, incremental indexing, parallel retrieval, adaptive model routing, and explicit service-level objectives."),
        ("Broader domain validation", "Test the workflow on additional grid protection, control panel, substation automation, transformer, switchgear, and utility procurement packages."),
    ]
    for idx, (name, detail) in enumerate(future, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{idx}. {name}")
        set_run(r, size=11.2, color=GREEN, bold=True)
        add_body(doc, detail)

    page_break(doc)
    doc.add_heading("Conclusion", level=1)
    add_body(
        doc,
        "GridSpec demonstrates a practical pattern for controlled agentic AI in engineering proposals. The strongest result is not the number of agents. It is the separation of responsibilities: deterministic code handles repeatable work, RAG supplies evidence, agents perform bounded reasoning, MCP controls tool access, and engineers approve the important outcomes.",
    )
    add_body(
        doc,
        "The prototype shows that a large RFQ can be converted into a traceable requirement story, routed through evidence-aware compliance, and assembled into a reviewable solution package. The next step is rigorous accuracy evaluation and production hardening so the under-one-hour initial-draft target can be measured and trusted across additional projects.",
    )

    # Apply header and footer to any sections added by Word internals.
    for section in doc.sections:
        if not section.header.paragraphs[0].text:
            add_header_footer(section)
        if len(section.footer.paragraphs) < 2:
            add_page_number(section)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
